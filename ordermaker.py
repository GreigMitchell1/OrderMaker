from docx import Document
from docx.shared import Inches, Pt
from docx.enum.style import WD_STYLE_TYPE

document = Document()

style=document.styles.add_style('indent', WD_STYLE_TYPE.PARAGRAPH)
paragraph_format = style.paragraph_format
paragraph_format.left_indent = Inches(0.1)
paragraph_format.first_line_indent = Inches(-0.25)
paragraph_format.space_before = Pt(12)
paragraph_format.window_control = True

document.add_heading('Document Title', 00)

p = document.add_paragraph('A test paragraph ')
p.add_run('bold text').bold = True
p.add_run(' and some ')
p.add_run('italic. ').italic = True

document.add_heading('Heading Level 1', level=1)
document.add_paragraph('Intense Quote', style='Intense Quote')

document.add_paragraph(
    'first item in unordered list', style='List Bullet'
)

document.add_paragraph(
    'first item in ordered list', style='List Number'
)

document.add_picture('clippy.jpg', width=Inches(1.25))

records = (
    (3, '101', 'Spam'),
    (7, '422', 'Eggs'),
    (4, '631', 'Spam, spam, eggs, and spam')
)

table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id' 
hdr_cells[2].text = 'Desc'
for qty, id, desc in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = id
    row_cells[2].text = desc

document.add_page_break()

document.save('ordertest2.docx')